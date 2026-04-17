"""Word document parser using python-docx."""

from __future__ import annotations

from pathlib import Path

from app.models.mcp import ParsedDocument


async def parse_docx(file_path: str | Path) -> ParsedDocument:
    """Parse a DOCX file and extract text."""
    try:
        from docx import Document

        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return ParsedDocument(
            source=str(file_path),
            content="\n\n".join(paragraphs),
            file_type=".docx",
            metadata={"paragraph_count": len(paragraphs)},
        )
    except ImportError:
        return ParsedDocument(source=str(file_path), content="Error: python-docx not installed", file_type=".docx")
