"""File processing pipeline — routes files to appropriate parsers.

Provides a unified entry point for processing any supported file format.
Automatically detects file type and dispatches to the correct parser.
"""

from __future__ import annotations

import time
from pathlib import Path
from typing import Any

import structlog

from app.config import settings
from app.models.mcp import FileProcessingRequest, FileProcessingResponse, ParsedDocument

logger = structlog.get_logger("file_processing.pipeline")

# Extension -> parser module mapping
PARSER_REGISTRY: dict[str, str] = {
    ".pdf": "app.file_processing.parsers.pdf_parser",
    ".docx": "app.file_processing.parsers.docx_parser",
    ".xlsx": "app.file_processing.parsers.xlsx_parser",
    ".pptx": "app.file_processing.parsers.pptx_parser",
    ".md": "app.file_processing.parsers.markdown_parser",
    ".txt": "app.file_processing.parsers.markdown_parser",
    ".py": "app.file_processing.parsers.code_parser",
    ".java": "app.file_processing.parsers.code_parser",
    ".ts": "app.file_processing.parsers.code_parser",
    ".js": "app.file_processing.parsers.code_parser",
    ".png": "app.file_processing.parsers.image_parser",
    ".jpg": "app.file_processing.parsers.image_parser",
    ".jpeg": "app.file_processing.parsers.image_parser",
}


class FileProcessingPipeline:
    """Unified file processing pipeline.

    Flow: Detect type -> Parse -> Extract tables -> Extract metadata -> Return
    """

    def __init__(self) -> None:
        self._parsers: dict[str, Any] = {}

    async def process(self, request: FileProcessingRequest) -> FileProcessingResponse:
        """Process a single file through the pipeline."""
        start = time.perf_counter()
        file_path = Path(request.file_path)

        if not file_path.exists():
            return FileProcessingResponse(
                file_path=request.file_path,
                content=f"Error: File not found: {request.file_path}",
                latency_ms=(time.perf_counter() - start) * 1000,
            )

        ext = file_path.suffix.lower()
        supported = settings.file_processing.supported_extensions
        if ext not in supported:
            return FileProcessingResponse(
                file_path=request.file_path,
                file_type=ext,
                content=f"Error: Unsupported file format: {ext}",
                latency_ms=(time.perf_counter() - start) * 1000,
            )

        try:
            # Dispatch to parser
            doc = await self._parse(file_path, ext)
            elapsed = (time.perf_counter() - start) * 1000

            await logger.ainfo(
                "file_processed",
                path=str(file_path),
                type=ext,
                content_len=len(doc.content),
                latency_ms=round(elapsed, 2),
            )

            return FileProcessingResponse(
                file_path=request.file_path,
                file_type=ext,
                content=doc.content,
                pages=doc.pages,
                tables=doc.tables,
                metadata=doc.metadata,
                latency_ms=elapsed,
            )
        except Exception as exc:
            elapsed = (time.perf_counter() - start) * 1000
            await logger.aerror("file_processing_error", path=str(file_path), error=str(exc))
            return FileProcessingResponse(
                file_path=request.file_path,
                file_type=ext,
                content=f"Error: {exc}",
                latency_ms=elapsed,
            )

    async def _parse(self, file_path: Path, ext: str) -> ParsedDocument:
        """Dispatch to the appropriate parser."""
        # Text-based formats
        if ext in (".md", ".txt", ".py", ".java", ".ts", ".js", ".json", ".yaml", ".yml", ".toml", ".csv"):
            content = file_path.read_text(encoding="utf-8", errors="replace")
            return ParsedDocument(source=str(file_path), content=content, file_type=ext)

        # PDF via PyMuPDF
        if ext == ".pdf":
            return await self._parse_pdf(file_path)

        # DOCX via python-docx
        if ext == ".docx":
            return await self._parse_docx(file_path)

        # XLSX via openpyxl
        if ext == ".xlsx":
            return await self._parse_xlsx(file_path)

        return ParsedDocument(
            source=str(file_path),
            content=f"Unsupported binary format: {ext}",
            file_type=ext,
        )

    async def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """Parse PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(str(file_path))
            pages_text = []
            for page in doc:
                pages_text.append(page.get_text())
            doc.close()

            return ParsedDocument(
                source=str(file_path),
                content="\n\n".join(pages_text),
                file_type=".pdf",
                pages=len(pages_text),
            )
        except ImportError:
            return ParsedDocument(
                source=str(file_path),
                content="Error: PyMuPDF not installed. Run: pip install PyMuPDF",
                file_type=".pdf",
            )

    async def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """Parse DOCX using python-docx."""
        try:
            from docx import Document

            doc = Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return ParsedDocument(
                source=str(file_path),
                content="\n\n".join(paragraphs),
                file_type=".docx",
            )
        except ImportError:
            return ParsedDocument(
                source=str(file_path),
                content="Error: python-docx not installed. Run: pip install python-docx",
                file_type=".docx",
            )

    async def _parse_xlsx(self, file_path: Path) -> ParsedDocument:
        """Parse XLSX using openpyxl."""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(str(file_path), read_only=True)
            content_parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                content_parts.append(f"## Sheet: {sheet_name}")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    content_parts.append(" | ".join(cells))
            wb.close()

            return ParsedDocument(
                source=str(file_path),
                content="\n".join(content_parts),
                file_type=".xlsx",
            )
        except ImportError:
            return ParsedDocument(
                source=str(file_path),
                content="Error: openpyxl not installed. Run: pip install openpyxl",
                file_type=".xlsx",
            )

    async def process_batch(self, requests: list[FileProcessingRequest]) -> list[FileProcessingResponse]:
        """Process multiple files."""
        import asyncio

        return await asyncio.gather(*[self.process(req) for req in requests])
